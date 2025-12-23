import argparse
import os
import shutil
import subprocess
from datetime import datetime
from pathlib import Path

from dotenv import load_dotenv


DEFAULT_TRANSCRIPT_MODEL_OPENAI = "gpt-4o-mini-transcribe"
DEFAULT_MINUTES_MODEL_OPENAI = "gpt-4o-mini"
DEFAULT_TRANSCRIPT_MODEL_GROQ = "whisper-large-v3"
DEFAULT_MINUTES_MODEL_GROQ = "llama-3.1-70b-versatile"
DEFAULT_TRANSCRIPT_MODEL_LOCAL = "openai/whisper-medium"
DEFAULT_MINUTES_MODEL_LOCAL = "meta-llama/Llama-3.2-3B-Instruct"


def run_cmd(args):
    result = subprocess.run(args, capture_output=True, text=True)
    if result.returncode != 0:
        raise RuntimeError(
            f"Command failed: {' '.join(args)}\n{result.stderr.strip()}"
        )
    return result.stdout.strip()


def ensure_dir(path):
    Path(path).mkdir(parents=True, exist_ok=True)


def check_gpu():
    info_lines = []
    try:
        import torch

        cuda_available = torch.cuda.is_available()
        info_lines.append(f"torch.cuda.is_available: {cuda_available}")
    except Exception as exc:
        return False, f"torch not available: {exc}"

    nvidia_smi = shutil.which("nvidia-smi")
    if nvidia_smi:
        try:
            output = run_cmd([nvidia_smi])
            info_lines.append(output)
        except Exception as exc:
            info_lines.append(f"nvidia-smi error: {exc}")
    else:
        info_lines.append("nvidia-smi not found")

    return cuda_available, "\n".join(info_lines)


def require_ffmpeg():
    if not shutil.which("ffmpeg") or not shutil.which("ffprobe"):
        raise RuntimeError("ffmpeg and ffprobe are required on PATH.")


def video_to_audio(video_path, audio_dir, force=False):
    require_ffmpeg()
    video_path = Path(video_path)
    audio_dir = Path(audio_dir)
    ensure_dir(audio_dir)
    audio_path = audio_dir / f"{video_path.stem}.mp3"

    if audio_path.exists() and not force:
        return audio_path

    run_cmd(
        [
            "ffmpeg",
            "-i",
            str(video_path),
            "-q:a",
            "0",
            "-map",
            "a",
            str(audio_path),
            "-y",
        ]
    )
    return audio_path


def audio_duration_minutes(audio_path):
    require_ffmpeg()
    result = run_cmd(
        [
            "ffprobe",
            "-v",
            "error",
            "-show_entries",
            "format=duration",
            "-of",
            "default=noprint_wrappers=1:nokey=1",
            str(audio_path),
        ]
    )
    duration_seconds = float(result)
    return duration_seconds / 60.0


def split_audio(audio_path, audio_dir, threshold_min, force=False):
    require_ffmpeg()
    audio_path = Path(audio_path)
    audio_dir = Path(audio_dir)
    ensure_dir(audio_dir)

    duration_min = audio_duration_minutes(audio_path)
    if duration_min <= threshold_min:
        return [audio_path]

    num_chunks = int(duration_min // threshold_min) + (
        1 if duration_min % threshold_min > 0 else 0
    )
    chunk_paths = []
    for i in range(num_chunks):
        start_time = i * threshold_min * 60
        chunk_path = audio_dir / f"{audio_path.stem}_part{i + 1}.mp3"
        chunk_paths.append(chunk_path)
        if chunk_path.exists() and not force:
            continue
        run_cmd(
            [
                "ffmpeg",
                "-i",
                str(audio_path),
                "-ss",
                str(start_time),
                "-t",
                str(threshold_min * 60),
                str(chunk_path),
                "-y",
            ]
        )
    return chunk_paths


def get_list_level(para):
    p_pr = para._p.pPr
    if p_pr is not None and p_pr.numPr is not None:
        ilvl = p_pr.numPr.ilvl
        if ilvl is not None:
            return int(ilvl.val)
    return None


def docx_to_markdown(docx_path):
    from docx import Document

    document = Document(docx_path)
    markdown_lines = []
    counters = [0, 0, 0, 0, 0]

    for para in document.paragraphs:
        raw_text = para.text.strip()
        if not raw_text:
            continue

        if para.style.name.startswith("Heading"):
            level = int(para.style.name.replace("Heading ", ""))
            markdown_lines.append("#" * level + " " + raw_text)
            continue

        list_level = get_list_level(para)
        if list_level is not None:
            counters[list_level] += 1
            for i in range(list_level + 1, len(counters)):
                counters[i] = 0
            number = ".".join(str(counters[i]) for i in range(list_level + 1))

            formatted_text = ""
            for run in para.runs:
                if run.bold:
                    formatted_text += f"**{run.text}**"
                else:
                    formatted_text += run.text
            if not formatted_text.strip():
                formatted_text = raw_text
            markdown_lines.append(f"{number}. {formatted_text.strip()}")
            continue

        if "List Bullet" in para.style.name:
            formatted_text = ""
            for run in para.runs:
                if run.bold:
                    formatted_text += f"**{run.text}**"
                else:
                    formatted_text += run.text
            if not formatted_text.strip():
                formatted_text = raw_text
            markdown_lines.append(f"- {formatted_text.strip()}")
            continue

        formatted_text = ""
        for run in para.runs:
            if run.bold:
                formatted_text += f"**{run.text}**"
            else:
                formatted_text += run.text
        if not formatted_text.strip():
            formatted_text = raw_text
        markdown_lines.append(formatted_text.strip())

    return "\n\n".join(markdown_lines)


def markdown_to_docx(markdown_text, output_path):
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    title = doc.add_heading("Compte-rendu de reunion", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    lines = markdown_text.split("\n")
    for line in lines:
        line = line.strip()
        if not line:
            continue

        if line.startswith("###"):
            doc.add_heading(line.replace("###", "").strip(), level=3)
        elif line.startswith("##"):
            doc.add_heading(line.replace("##", "").strip(), level=2)
        elif line.startswith("#"):
            doc.add_heading(line.replace("#", "").strip(), level=1)
        elif line.startswith("**") and line.endswith("**"):
            doc.add_heading(line.replace("**", "").strip(), level=3)
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:], style="List Bullet")
        else:
            doc.add_paragraph(line)

    ensure_dir(Path(output_path).parent)
    doc.save(output_path)


def load_agenda_markdown(agenda_path, output_dir, force=False):
    if not agenda_path:
        return None
    agenda_path = Path(agenda_path)
    output_dir = Path(output_dir)
    ensure_dir(output_dir)
    output_path = output_dir / f"{agenda_path.stem}.md"

    if output_path.exists() and not force:
        return output_path.read_text(encoding="utf-8")

    markdown_text = docx_to_markdown(agenda_path)
    output_path.write_text(markdown_text, encoding="utf-8")
    return markdown_text


def build_messages(full_text, agenda_markdown):
    current_date = datetime.now().strftime("%y-%m-%d")
    system_message = (
        f"Current date: {current_date}.\n"
        "You produce minutes of meetings from transcripts, with summary, key "
        "discussion points, takeaways, votes, and action items with board "
        "members and general manager, in markdown format without code blocks, "
        "in french."
    )

    user_prompt = (
        "Below is an extract transcript of the C.A. of the Jeune Chambre de "
        "Drummond.\n"
        "Please write minutes in markdown without code blocks, including:\n"
        "- a summary with attendees, date, location\n"
        "- discussion points\n"
        "- takeaways\n"
        "- votes\n"
        "- action items with board members and general manager.\n"
        "If a section is not applicable, say so.\n"
        "If a vote is taken, please write it in the format:\n"
        "**Resolution [current_date yy-mm-dd]-[incremental_number]**\n"
        "Il est propose par [name_1] et appuye par [name_2] de "
        "[description of the motion].\n"
        "**-Adopte a l'unanimite-**\n"
        "Example:\n"
        "**Resolution 25-12-31-01**\n"
        "Il est propose par Frederic et appuye par Laurianne d'ouvrir la "
        "reunion a 7h03. Le quorum est constate.\n"
        "**-Adopte a l'unanimite-**\n"
        "======\n"
    )

    if agenda_markdown and agenda_markdown.strip():
        user_prompt += (
            "Use the agenda to help structure the minutes. Here is the agenda:\n"
            f"{agenda_markdown}\n"
        )

    user_prompt += f"Here is the transcript:\n{full_text}\n"

    return [
        {"role": "system", "content": system_message},
        {"role": "user", "content": user_prompt},
    ]


def transcribe_chunks(chunks, provider, model, language, output_dir, force=False):
    output_dir = Path(output_dir)
    ensure_dir(output_dir)

    transcripts = []
    transcript_paths = []

    if provider == "local":
        available, info = check_gpu()
        print(info)
        if not available:
            print("Warning: GPU not available for local transcription.")
        from transformers import pipeline
        import torch

        device = "cuda" if torch.cuda.is_available() else "cpu"
        pipe = pipeline(
            "automatic-speech-recognition",
            model=model,
            torch_dtype=torch.float16 if device == "cuda" else torch.float32,
            device=device,
            return_timestamps=True,
        )

        for chunk_path in chunks:
            transcript_path = output_dir / f"{Path(chunk_path).stem}.txt"
            if transcript_path.exists() and not force:
                text = transcript_path.read_text(encoding="utf-8")
            else:
                result = pipe(str(chunk_path), generate_kwargs={"language": language})
                text = result["text"]
                transcript_path.write_text(text, encoding="utf-8")
            transcripts.append(text)
            transcript_paths.append(transcript_path)

    else:
        from openai import OpenAI

        if provider == "groq":
            api_key = os.getenv("GROQ_API_KEY")
            client = OpenAI(api_key=api_key, base_url="https://api.groq.com/openai/v1")
        else:
            client = OpenAI()

        for chunk_path in chunks:
            transcript_path = output_dir / f"{Path(chunk_path).stem}.txt"
            if transcript_path.exists() and not force:
                text = transcript_path.read_text(encoding="utf-8")
            else:
                with open(chunk_path, "rb") as audio_file:
                    text = client.audio.transcriptions.create(
                        model=model,
                        file=audio_file,
                        response_format="text",
                        language=language,
                    )
                transcript_path.write_text(text, encoding="utf-8")
            transcripts.append(text)
            transcript_paths.append(transcript_path)

    return transcripts, transcript_paths


def merge_transcripts(transcripts, output_path, force=False):
    output_path = Path(output_path)
    if output_path.exists() and not force:
        return output_path.read_text(encoding="utf-8")

    full_text = "".join(transcripts)
    ensure_dir(output_path.parent)
    output_path.write_text(full_text, encoding="utf-8")
    return full_text


def generate_minutes(messages, provider, model, max_new_tokens):
    if provider == "local":
        available, info = check_gpu()
        print(info)
        if not available:
            print("Warning: GPU not available for local generation.")
        from transformers import AutoModelForCausalLM, AutoTokenizer, BitsAndBytesConfig
        import torch

        quant_config = None
        if torch.cuda.is_available():
            quant_config = BitsAndBytesConfig(
                load_in_4bit=True,
                bnb_4bit_use_double_quant=True,
                bnb_4bit_compute_dtype=torch.bfloat16,
                bnb_4bit_quant_type="nf4",
            )

        tokenizer = AutoTokenizer.from_pretrained(model)
        tokenizer.pad_token = tokenizer.eos_token
        inputs = tokenizer.apply_chat_template(messages, return_tensors="pt")
        if torch.cuda.is_available():
            inputs = inputs.to("cuda")
        model_obj = AutoModelForCausalLM.from_pretrained(
            model,
            device_map="auto",
            quantization_config=quant_config,
        )
        outputs = model_obj.generate(inputs, max_new_tokens=max_new_tokens)
        generated_text = tokenizer.decode(
            outputs[0][len(inputs[0]) :], skip_special_tokens=True
        )
        if generated_text.startswith("assistant"):
            generated_text = generated_text.replace("assistant", "", 1).strip()
        return generated_text

    from openai import OpenAI

    if provider == "groq":
        api_key = os.getenv("GROQ_API_KEY")
        client = OpenAI(api_key=api_key, base_url="https://api.groq.com/openai/v1")
    else:
        client = OpenAI()

    response = client.chat.completions.create(model=model, messages=messages)
    return response.choices[0].message.content


def maybe_login_hf():
    hf_token = os.getenv("HF_TOKEN")
    if not hf_token:
        return
    try:
        from huggingface_hub import login

        login(hf_token, add_to_git_credential=True)
    except Exception as exc:
        print(f"Warning: HuggingFace login failed: {exc}")


def resolve_models(provider, transcript_model, minutes_model):
    if provider == "openai":
        t_default = DEFAULT_TRANSCRIPT_MODEL_OPENAI
        m_default = DEFAULT_MINUTES_MODEL_OPENAI
    elif provider == "groq":
        t_default = DEFAULT_TRANSCRIPT_MODEL_GROQ
        m_default = DEFAULT_MINUTES_MODEL_GROQ
    else:
        t_default = DEFAULT_TRANSCRIPT_MODEL_LOCAL
        m_default = DEFAULT_MINUTES_MODEL_LOCAL
    return transcript_model or t_default, minutes_model or m_default


def generate_command(args):
    if not args.video and not args.audio:
        raise ValueError("Provide --video or --audio.")
    if args.video and args.audio:
        raise ValueError("Provide only one of --video or --audio.")

    ensure_dir("audio")
    ensure_dir("transcripts")
    ensure_dir("minutes/markdown")
    ensure_dir("minutes/word")
    ensure_dir("agendas/markdown")

    if args.video:
        audio_path = video_to_audio(args.video, "audio", force=args.force)
    else:
        audio_path = Path(args.audio)

    if not audio_path.exists():
        raise FileNotFoundError(f"Audio file not found: {audio_path}")

    chunks = split_audio(
        audio_path, "audio", threshold_min=args.split_threshold_min, force=args.force
    )

    transcript_model, minutes_model = resolve_models(
        args.provider, args.transcript_model, args.minutes_model
    )

    if args.provider == "local":
        maybe_login_hf()

    transcripts, transcript_paths = transcribe_chunks(
        chunks,
        provider=args.provider,
        model=transcript_model,
        language=args.language,
        output_dir="transcripts",
        force=args.force,
    )

    base_name = Path(audio_path).stem
    merged_path = Path("transcripts") / f"{base_name}_full.txt"
    full_text = merge_transcripts(transcripts, merged_path, force=args.force)

    agenda_markdown = load_agenda_markdown(
        args.agenda, "agendas/markdown", force=args.force
    )

    messages = build_messages(full_text, agenda_markdown)

    minutes_md_path = Path("minutes/markdown") / f"{base_name}.md"
    if minutes_md_path.exists() and not args.force:
        minutes_text = minutes_md_path.read_text(encoding="utf-8")
    else:
        minutes_text = generate_minutes(
            messages,
            provider=args.provider,
            model=minutes_model,
            max_new_tokens=args.max_new_tokens,
        )
        minutes_md_path.write_text(minutes_text, encoding="utf-8")

    if not args.skip_word:
        minutes_docx_path = Path("minutes/word") / f"{base_name}.docx"
        if not minutes_docx_path.exists() or args.force:
            markdown_to_docx(minutes_text, minutes_docx_path)

    print(f"Audio: {audio_path}")
    print(f"Chunks: {len(chunks)}")
    print(f"Transcripts: {[str(p) for p in transcript_paths]}")
    print(f"Full transcript: {merged_path}")
    if agenda_markdown:
        print("Agenda markdown saved in agendas/markdown")
    print(f"Minutes markdown: {minutes_md_path}")
    if not args.skip_word:
        print(f"Minutes docx: {minutes_docx_path}")


def docx_to_md_command(args):
    text = docx_to_markdown(args.input)
    output_path = Path(args.output) if args.output else Path(args.input).with_suffix(".md")
    output_path.write_text(text, encoding="utf-8")
    print(f"Saved: {output_path}")


def md_to_docx_command(args):
    text = Path(args.input).read_text(encoding="utf-8")
    output_path = Path(args.output) if args.output else Path(args.input).with_suffix(".docx")
    markdown_to_docx(text, output_path)
    print(f"Saved: {output_path}")


def check_gpu_command(_args):
    available, info = check_gpu()
    print(info)
    if available:
        print("GPU available")
    else:
        print("GPU not available")


def build_parser():
    parser = argparse.ArgumentParser(
        description="Generate meeting minutes from audio or video."
    )
    subparsers = parser.add_subparsers(dest="command", required=True)

    gen = subparsers.add_parser("generate", help="Generate minutes from audio/video.")
    gen.add_argument("--video", type=str, help="Path to a video file.")
    gen.add_argument("--audio", type=str, help="Path to an audio file.")
    gen.add_argument(
        "--provider",
        choices=["openai", "groq", "local"],
        default="openai",
        help="Provider for transcription and minutes generation.",
    )
    gen.add_argument(
        "--transcript-model", type=str, help="Model for transcription (optional)."
    )
    gen.add_argument(
        "--minutes-model", type=str, help="Model for minutes generation (optional)."
    )
    gen.add_argument(
        "--agenda", type=str, help="Optional agenda docx to use as template."
    )
    gen.add_argument(
        "--split-threshold-min",
        type=int,
        default=10,
        help="Split audio into chunks above this duration (minutes).",
    )
    gen.add_argument(
        "--language",
        type=str,
        default="fr",
        help="Language code for transcription.",
    )
    gen.add_argument(
        "--max-new-tokens",
        type=int,
        default=2000,
        help="Max new tokens for local generation.",
    )
    gen.add_argument("--force", action="store_true", help="Recompute all steps.")
    gen.add_argument("--skip-word", action="store_true", help="Skip Word export.")

    docx_md = subparsers.add_parser("docx-to-md", help="Convert docx to markdown.")
    docx_md.add_argument("input", type=str, help="Path to docx file.")
    docx_md.add_argument("--output", type=str, help="Output markdown path.")

    md_docx = subparsers.add_parser("md-to-docx", help="Convert markdown to docx.")
    md_docx.add_argument("input", type=str, help="Path to markdown file.")
    md_docx.add_argument("--output", type=str, help="Output docx path.")

    gpu = subparsers.add_parser("check-gpu", help="Check GPU availability.")

    gen.set_defaults(func=generate_command)
    docx_md.set_defaults(func=docx_to_md_command)
    md_docx.set_defaults(func=md_to_docx_command)
    gpu.set_defaults(func=check_gpu_command)

    return parser


def main():
    load_dotenv()
    parser = build_parser()
    args = parser.parse_args()
    args.func(args)


if __name__ == "__main__":
    main()
